"""Resolve a resume from a PDF upload, a Google Drive link, or an external link,
and return clean plain text for analysis.
"""

from __future__ import annotations

import io
import re

import httpx
from bs4 import BeautifulSoup
from fastapi import HTTPException

_DRIVE_FILE_RE = re.compile(r"/file/d/([a-zA-Z0-9_-]+)")
_DRIVE_ID_RE = re.compile(r"[?&]id=([a-zA-Z0-9_-]+)")

_HEADERS = {"User-Agent": "Mozilla/5.0 (ResumeEnhancer ingestion bot)"}


def _normalize_drive_url(url: str) -> str:
    """Convert a shareable Drive URL into a direct-download URL."""
    m = _DRIVE_FILE_RE.search(url) or _DRIVE_ID_RE.search(url)
    if not m:
        raise HTTPException(status_code=400, detail="Could not parse Google Drive file id from URL")
    return f"https://drive.google.com/uc?export=download&id={m.group(1)}"


def _append_links(text: str, urls: list[str]) -> str:
    """Append real hyperlink URLs so the model & link checks see them even when the
    visible anchor text is just 'Link' / 'GitHub' / 'LeetCode'."""
    uniq = [u for u in dict.fromkeys(u.strip() for u in urls if u) if u.lower().startswith(("http://", "https://"))]
    if not uniq:
        return text
    return text + "\n\nHYPERLINKS DETECTED IN RESUME:\n" + "\n".join(uniq)


def _pdf_page_links(page) -> list[str]:
    urls: list[str] = []
    try:
        for h in (page.hyperlinks or []):
            if h.get("uri"):
                urls.append(h["uri"])
    except Exception:
        pass
    try:
        for a in (page.annots or []):
            data = a.get("data") or {}
            uri = data.get("uri") or data.get("URI")
            action = data.get("A") if isinstance(data, dict) else None
            if not uri and isinstance(action, dict):
                uri = action.get("URI")
            if isinstance(uri, bytes):
                uri = uri.decode("utf-8", errors="ignore")
            if uri:
                urls.append(str(uri))
    except Exception:
        pass
    return urls


def extract_text_from_pdf_bytes(data: bytes) -> str:
    import pdfplumber

    text_parts: list[str] = []
    urls: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
            urls.extend(_pdf_page_links(page))
    return _append_links(_clean("\n".join(text_parts)), urls)


def extract_text_from_docx_bytes(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs)
    urls: list[str] = []
    try:
        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype.lower():
                urls.append(rel.target_ref)
    except Exception:
        pass
    return _append_links(_clean(text), urls)


def _is_white(color) -> bool:
    """True if a pdfplumber fill color is white / near-white (hidden text)."""
    if color is None:
        return False
    try:
        if isinstance(color, (int, float)):
            return color >= 0.95
        comps = list(color)
        return bool(comps) and all(float(c) >= 0.95 for c in comps)
    except (TypeError, ValueError):
        return False


def pdf_meta(data: bytes) -> dict:
    """Page count, embedded-image count (possible photo), and hidden/white/tiny text
    used for keyword-stuffing (anti-gaming)."""
    import pdfplumber

    pages = 0
    images = 0
    hidden: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                images += len(page.images or [])
                suspicious = []
                for ch in page.chars:
                    size = ch.get("size") or 10
                    if _is_white(ch.get("non_stroking_color")) or size < 3:
                        suspicious.append(ch.get("text", ""))
                snippet = "".join(suspicious).strip()
                if len(snippet) > 8:
                    hidden.append(snippet[:200])
    except Exception:
        pass
    return {"pages": pages, "images": images, "hidden_text": hidden}


def docx_meta(data: bytes) -> dict:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
        images = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    except Exception:
        images = 0
    return {"pages": None, "images": images, "hidden_text": []}


def extract_text_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    urls = [a.get("href") for a in soup.find_all("a") if a.get("href")]
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return _append_links(_clean(soup.get_text(separator="\n")), urls)


def _clean(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    return "\n".join(lines).strip()


def _fetch(url: str) -> httpx.Response:
    try:
        with httpx.Client(follow_redirects=True, timeout=30, headers=_HEADERS) as client:
            resp = client.get(url)
            resp.raise_for_status()
            return resp
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=400, detail=f"Failed to fetch resume URL: {exc}") from exc


_EMPTY_META = {"pages": None, "images": 0, "hidden_text": []}


def _bytes_to_text(content: bytes, content_type: str, url: str) -> tuple[str, dict]:
    ct = content_type.lower()
    if "pdf" in ct or url.lower().endswith(".pdf") or content[:4] == b"%PDF":
        return extract_text_from_pdf_bytes(content), pdf_meta(content)
    if "word" in ct or url.lower().endswith(".docx"):
        return extract_text_from_docx_bytes(content), docx_meta(content)
    if "html" in ct or content[:15].lstrip().lower().startswith(b"<!doctype") or b"<html" in content[:512].lower():
        return extract_text_from_html(content.decode("utf-8", errors="ignore")), dict(_EMPTY_META)
    # Fallback: treat as plain text
    return _clean(content.decode("utf-8", errors="ignore")), dict(_EMPTY_META)


def text_from_upload(filename: str, data: bytes) -> tuple[str, dict]:
    name = filename.lower()
    if name.endswith(".pdf") or data[:4] == b"%PDF":
        return extract_text_from_pdf_bytes(data), pdf_meta(data)
    if name.endswith(".docx"):
        return extract_text_from_docx_bytes(data), docx_meta(data)
    raise HTTPException(status_code=400, detail="Unsupported file type. Upload a PDF or DOCX.")


def text_from_drive(url: str) -> tuple[str, dict]:
    direct = _normalize_drive_url(url)
    resp = _fetch(direct)
    return _bytes_to_text(resp.content, resp.headers.get("content-type", ""), direct)


def text_from_link(url: str) -> tuple[str, dict]:
    resp = _fetch(url)
    text, meta = _bytes_to_text(resp.content, resp.headers.get("content-type", ""), url)
    if not text:
        raise HTTPException(status_code=400, detail="No readable content found at the link")
    return text, meta
